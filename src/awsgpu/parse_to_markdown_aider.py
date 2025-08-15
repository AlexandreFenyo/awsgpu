#!/usr/bin/env python3

"""
parse_to_markdown_aider.py

Script (FR) pour convertir un fichier Word (.docx) en Markdown en utilisant:
- le module awsgpu.parse_word.parse_docx_to_blocks pour extraire le contenu en "blocks"
- (optionnel) le module markitdown si présent
- l'API OpenAI (modèle gpt-4o) pour générer des textes alternatifs (alt text) pour les images
  La clé d'API doit être fournie via la variable d'environnement OPENAIAPIKEY.

Fonctionnement:
- Extrait le texte et les blocs via parse_docx_to_blocks()
- Extrait les images du fichier .docx et les sauvegarde dans un dossier <output_basename>_images/
- Pour chaque image, demande à l'API OpenAI (gpt-4o) une courte description (alt) en français
- Génère un fichier Markdown avec le texte converti et une section listant les images avec leur alt text

Remarques:
- Le script essaie d'utiliser le nouveau client "openai.OpenAI" si disponible,
  sinon il essaie la bibliothèque "openai" traditionnelle.
- Si le module markitdown est présent, il l'utilisera si possible. Sinon, une conversion simple est appliquée.
- Ne modifie aucun autre fichier du dépôt.

Commandes d'exemple (à lancer depuis la racine du projet):
# python3 -m src.awsgpu.parse_to_markdown_aider input.docx output.md
# python3 -m src.awsgpu.parse_to_markdown_aider "Mon Document.docx" out.md
"""

from __future__ import annotations

import base64
import os
import sys
import typing as t
from pathlib import Path

try:
    from docx import Document
except Exception as e:
    raise RuntimeError(
        "python-docx is required. Install with: pip install python-docx"
    ) from e

# Using MarkItDown exclusively; parse_word is not used.
# parse_docx_to_blocks removed.

# MarkItDown is required
_has_markitdown = True
_markitdown_converter = None
try:
    import markitdown  # type: ignore
    if hasattr(markitdown, "MarkItDown"):
        _markitdown_converter = markitdown.MarkItDown()  # type: ignore
    else:
        _markitdown_converter = markitdown  # type: ignore
except Exception as e:
    raise RuntimeError("markitdown is required. Install with: pip install markitdown") from e


def _get_openai_api_key() -> str:
    key = os.environ.get("OPENAIAPIKEY")
    if not key:
        raise RuntimeError(
            "Environment variable OPENAIAPIKEY is not set. Please export your OpenAI API key."
        )
    return key


def _describe_image_with_openai(image_bytes: bytes, api_key: str) -> str:
    """
    Ask OpenAI (model gpt-4o) for a short French alt-text for the image.
    The function tries to use the newer OpenAI client if available, otherwise
    falls back to the classic openai package. It is robust to failures and
    returns a fallback short description on error.
    """
    b64 = base64.b64encode(image_bytes).decode("ascii")

    prompt = (
        "Voici une image encodée en base64:\n\n"
        f"{b64}\n\n"
        "Donne une courte description en français (texte alt) adaptée pour Markdown.\n"
        "Répond uniquement par la description (une phrase courte, max ~12 mots)."
    )

    # Try new OpenAI client first
    try:
        try:
            # new client
            from openai import OpenAI  # type: ignore

            client = OpenAI(api_key=api_key)  # type: ignore
            # responses.create returns different shapes; try to extract text robustly
            resp = client.responses.create(model="gpt-4o", input=prompt)  # type: ignore
            # Try common access patterns
            text = None
            if hasattr(resp, "output_text"):
                text = getattr(resp, "output_text")
            elif hasattr(resp, "output"):
                out = getattr(resp, "output")
                if isinstance(out, list) and out:
                    first = out[0]
                    if isinstance(first, dict) and "content" in first:
                        # content may be a list
                        content = first["content"]
                        if isinstance(content, list):
                            # find first output_text
                            for c in content:
                                if isinstance(c, dict) and c.get("type") in (
                                    "output_text",
                                    "message",
                                ):
                                    text = c.get("text") or c.get("content")
                                    if isinstance(text, list):
                                        text = " ".join(text)
                                    break
            if not text:
                # last resort: string conversion
                text = str(resp)
            return text.strip().strip('"').strip("'")
        except Exception:
            # fallback to classic openai package
            import openai  # type: ignore

            openai.api_key = api_key
            resp = openai.ChatCompletion.create(
                model="gpt-4o", messages=[{"role": "user", "content": prompt}], max_tokens=60
            )
            if "choices" in resp and resp["choices"]:
                content = resp["choices"][0].get("message", {}).get("content")
                if content:
                    return content.strip()
            return "image"
    except Exception:
        return "image"


def _extract_images_from_docx(path: str) -> list[tuple[str, bytes]]:
    """
    Extract images from a docx file.

    Returns a list of tuples: (mime_type, bytes)
    """
    doc = Document(path)
    images: list[tuple[str, bytes]] = []
    # doc.part.rels contains relationships; image parts have content_type starting with 'image/'
    for rel in doc.part.rels.values():
        try:
            target = rel.target_part
        except Exception:
            continue
        ctype = getattr(target, "content_type", "")
        blob = getattr(target, "blob", None)
        if ctype and ctype.startswith("image/") and blob:
            images.append((ctype, blob))
    return images


def _blocks_to_markdown(blocks: list[tuple[str, str, int]]) -> str:
    """
    Convert blocks (style, text, level) to a basic markdown representation.
    If markitdown is available and provides a useful API, prefer it.
    """
    if _has_markitdown and _markitdown_converter is not None:
        try:
            # Try some common markitdown APIs
            if hasattr(_markitdown_converter, "convert_docx"):
                # hypothetical convenience method
                return _markitdown_converter.convert_docx  # type: ignore
            if hasattr(_markitdown_converter, "convert"):
                # If convert accepts a list of blocks or text, try to join texts
                texts = "\n\n".join(b[1] for b in blocks)
                try:
                    return _markitdown_converter.convert(texts)  # type: ignore
                except Exception:
                    # If convert is not suitable, fall back
                    pass
        except Exception:
            # fallback to manual below
            pass

    lines: list[str] = []
    for style, text, level in blocks:
        st = style.lower() if isinstance(style, str) else ""
        if st.startswith("heading") or st.startswith("titre") or st.startswith("h"):
            # handle heading with level
            lvl = level or 1
            lvl = max(1, min(6, lvl))
            lines.append("#" * lvl + " " + text.strip())
        elif "list" in st or "bullet" in st:
            indent = "  " * max(0, level - 1)
            for line in text.splitlines():
                if line.strip():
                    lines.append(f"{indent}- {line.strip()}")
        elif "table" in st:
            # a very simple representation: include the text block as-is
            lines.append("\n".join(line.rstrip() for line in text.splitlines()))
        else:
            lines.append(text.strip())
        lines.append("")  # blank line after each block
    return "\n".join(lines).rstrip() + "\n"


def convert_docx_to_markdown_with_images(input_path: str, output_md: str) -> None:
    input_path = str(input_path)
    output_md = str(output_md)
    api_key = _get_openai_api_key()

    # 1) convert DOCX to Markdown using MarkItDown
    try:
        if hasattr(_markitdown_converter, "convert_docx"):
            md_text = _markitdown_converter.convert_docx(input_path)  # type: ignore
        else:
            # Fallback: extract plain text and pass to convert()
            doc = Document(input_path)
            texts = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            md_text = _markitdown_converter.convert(texts)  # type: ignore
    except Exception as e:
        raise RuntimeError("MarkItDown failed to convert DOCX to Markdown") from e

    # 3) extract images
    images = _extract_images_from_docx(input_path)

    out_dir = Path(output_md).resolve().parent
    base = Path(output_md).stem
    images_dir_name = f"{base}_images"
    images_dir = out_dir / images_dir_name
    images_dir.mkdir(parents=True, exist_ok=True)

    image_markdowns: list[str] = []
    for i, (mime, blob) in enumerate(images, start=1):
        ext = mime.split("/")[-1].split("+")[0]  # handle image/svg+xml
        filename = f"image_{i}.{ext}"
        fpath = images_dir / filename
        with open(fpath, "wb") as fh:
            fh.write(blob)

        # ask OpenAI for a short alt text (in French)
        try:
            alt = _describe_image_with_openai(blob, api_key)
        except Exception:
            alt = "image"

        # relative path from output_md to image
        rel_path = os.path.join(images_dir_name, filename)
        image_markdowns.append(f"![{alt}]({rel_path})")

    # Append images section if any images were found
    if image_markdowns:
        md_text = md_text.rstrip() + "\n\n---\n\n## Images extraites\n\n"
        md_text += "\n\n".join(image_markdowns) + "\n"

    # 4) write output
    with open(output_md, "w", encoding="utf-8") as out:
        out.write(md_text)

    print(f"Markdown saved to: {output_md}")
    if image_markdowns:
        print(f"Images saved to directory: {images_dir}")


def _usage_and_exit() -> None:
    prog = Path(sys.argv[0]).name
    print("Usage:")
    print(f"  python3 -m src.awsgpu.parse_to_markdown_aider input.docx output.md")
    print("")
    print("Example:")
    print(f"  python3 -m src.awsgpu.parse_to_markdown_aider 'Mon document.docx' output.md")
    sys.exit(2)


def main(argv: t.Optional[list[str]] = None) -> int:
    if argv is None:
        argv = sys.argv[1:]
    if len(argv) != 2:
        _usage_and_exit()
    input_path, output_md = argv
    convert_docx_to_markdown_with_images(input_path, output_md)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
